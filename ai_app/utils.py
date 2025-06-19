import os
import re
import json
import time

import requests
from decouple import config
from PyPDF2 import PdfReader
from docx import Document
import tempfile
from .common import log_execution_time, CommonUtils
import openai
from pdf2image import convert_from_path
import zipfile
import xml.etree.ElementTree as ET

def get_access_token():
    """
    Generate an access token using client credentials flow.
    """
    tenant_id = config("TENANT_ID")
    url = f"https://login.microsoftonline.com/{tenant_id}/oauth2/v2.0/token"
    data = {
        'grant_type': 'client_credentials',
        'client_id': config("CLIENT_ID"),
        'client_secret':config("CLIENT_SECRET"),
        'scope': 'https://graph.microsoft.com/.default',
    }
    response = requests.post(url, data=data)
    if response.status_code == 200:
        return response.json().get('access_token')
    else:
        raise Exception(f"Unable to get access token \n Error: {response.text}")


def upload_file_to_sharepoint(site_id, drive_id, folder_path, file_name, file_content):
    """
    Upload a file to SharePoint in the specified folder.
    """
    try:
        access_token = get_access_token()
        url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives/{drive_id}/root:/{folder_path}/{file_name}:/content"
        headers = {
            'Authorization': f'Bearer {access_token}',
            'Content-Type': 'application/octet-stream',
        }

        response = requests.put(url, headers=headers, data=file_content)
        if response.status_code in [200, 201]:
            return response.json()
        else:
            raise Exception(f"Failed to upload file: {response.json()}")
    except Exception as e:
        raise Exception(f"Error uploading file: {str(e)}")



def get_file_from_sharepoint(site_id, file_path, access_token):
    """
    Fetches a specific file from a predefined SharePoint path.
    """
    url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drive/root:/{file_path}"
    headers = {"Authorization": f"Bearer {access_token}"}
    response = requests.get(url, headers=headers)
    if response.status_code == 200:
        return response.json()
    else:
        raise Exception(f"Failed to fetch file: {response.json()}")


def get_file_by_project_id(site_id, library_path, project_id, access_token):
    """
    Fetches a file from a SharePoint document library where the project_id matches a file's metadata or name.
    """
    # Microsoft Graph API URL to list items in the library
    url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drive/root:/{library_path}:/children"
    headers = {"Authorization": f"Bearer {access_token}"}

    # Fetch all items in the document library
    response = requests.get(url, headers=headers)
    if response.status_code != 200:
        raise Exception(f"Failed to fetch files: {response.json()}")

    # Filter files based on the project_id in the metadata or file name
    items = response.json().get('value', [])
    for item in items:
        if item.get('file') and item.get('name') and project_id in item.get('name'):
            return item

    # If no matching file is found
    raise Exception(f"No file found for project ID: {project_id}")



def extract_text_from_pdf(file_path):
    """Extract text from a PDF file."""
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text


def extract_text_from_docx(file_path):
    """Extract text from a Word document."""
    doc = Document(file_path)
    text = "\n".join([p.text for p in doc.paragraphs])
    return text


def read_and_parse_documents(folder_path):
    """Read all PDF and DOCX files from the folder and return concatenated text."""
    all_text = ""
    discovery_questionnaire_text = None

    for file_name in os.listdir(folder_path):
        file_path = os.path.join(folder_path, file_name)
        if file_name.lower().endswith(".pdf"):
            text = extract_text_from_pdf(file_path)
        elif file_name.lower().endswith(".docx"):
            text = extract_text_from_docx(file_path)
        else:
            continue  # Skip non-PDF and non-DOCX files

        # Check for the Discovery Questionnaire document
        if "discovery questionnaire" in file_name.lower():
            discovery_questionnaire_text = text
        else:
            all_text += text + "\n"

    return all_text, discovery_questionnaire_text


def upload_questionnaire_to_sharepoint(file_path, project_id):
    try:
        access_token = get_access_token()
        headers = {
            "Authorization": f"Bearer {access_token}",
            "Content-Type": "application/json"
        }

        site_id = config("SITE_ID")
        discovery_drive = config("DISCOVERY_DRIVE")
        # Upload the file
        upload_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives/{discovery_drive}/root:/Discovery Questionnaire-{project_id}.docx:/content"

        with open(file_path, "rb") as file:
            response = requests.put(upload_url, headers=headers, data=file)

        if response.status_code not in [200, 201]:
            raise Exception(f"Failed to upload file: {response.json()}")

        # Extract the uploaded file's item ID
        item_id = response.json().get("id")

        # Get existing columns
        columns_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives/{discovery_drive}/items/{item_id}/listItem/fields"
        fields_response = requests.get(columns_url, headers=headers)

        if fields_response.status_code != 200:
            raise Exception(f"Failed to fetch columns: {fields_response.json()}")

        updated_project_id = {
            "ProjectId":project_id
        }

        # Update the columns
        update_response = requests.patch(columns_url, headers=headers, json=updated_project_id)

        if update_response.status_code != 200:
            raise Exception(f"Failed to update project_id: {update_response.json()}")

    except Exception as e:
        raise Exception(f"Error during SharePoint upload or update: {str(e)}")


def update_current_step(project_id, current_step, key="CurrentStep"):
    """
    Updates the CurrentStep field in the Project list for the specified project_id.
    """
    try:
        access_token = get_access_token()
        headers = {
            "Authorization": f"Bearer {access_token}",
            "Content-Type": "application/json"
        }

        site_id = config("SITE_ID")
        project_list_id = config("PROJECT_LIST")
        # Update URL for CurrentStep
        update_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/lists/{project_list_id}/items/{project_id}/fields"
        update_body = {key: current_step}

        # PATCH request to update CurrentStep
        response = requests.patch(update_url, headers=headers, json=update_body)

        if response.status_code != 200:
            raise Exception(f"Failed to update CurrentStep: {response.json()}")

    except Exception as e:
        raise Exception(f"Error updating CurrentStep: {str(e)}")


def upload_sow_to_sharepoint(file_path, project_id):
    try:
        access_token = get_access_token()
        headers = {
            "Authorization": f"Bearer {access_token}",
            "Content-Type": "application/json"
        }

        site_id = config("SITE_ID")
        sow_drive = config("SOW_DRIVE")
        # Upload the file
        upload_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives/{sow_drive}/root:/SOW-{project_id}.docx:/content"

        with open(file_path, "rb") as file:
            response = requests.put(upload_url, headers=headers, data=file)

        if response.status_code not in [200, 201]:
            raise Exception(f"Failed to upload file: {response.json()}")

        # Extract the uploaded file's item ID
        item_id = response.json().get("id")

        # Get existing columns
        columns_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives/{sow_drive}/items/{item_id}/listItem/fields"
        fields_response = requests.get(columns_url, headers=headers)

        if fields_response.status_code != 200:
            raise Exception(f"Failed to fetch columns: {fields_response.json()}")

        updated_project_id = {
            "ProjectId":project_id
        }

        # Update the columns
        update_response = requests.patch(columns_url, headers=headers, json=updated_project_id)

        if update_response.status_code != 200:
            raise Exception(f"Failed to update project_id: {update_response.json()}")

    except Exception as e:
        raise Exception(f"Error during SharePoint upload or update: {str(e)}")

def get_sharepoint_items(access_token, drive_url):
    """Fetch items from the SharePoint drive URL."""
    headers = {"Authorization": f"Bearer {access_token}"}
    response = requests.get(drive_url, headers=headers)
    response.raise_for_status()
    return response.json()

def get_taxonomy_item_id(access_token, items):
    headers = {"Authorization": f"Bearer {access_token}",}
    item_ids = [i["id"] for i in items["value"]]
    download_urls = [i["@microsoft.graph.downloadUrl"] for i in items["value"]]

    TAXONOMY_DRIVE_ID = config("TAXONOMY_DRIVE_ID")
    for ind, item_id in enumerate(item_ids):
        url = f"https://graph.microsoft.com/v1.0/drives/{TAXONOMY_DRIVE_ID}/items/{item_id}/listItem/fields"
        response = requests.get(url, headers=headers)
        values = response.json()
        if "isParsed" in values:
            if values["isParsed"] == False:
                return item_id, download_urls[ind]
    return -1, ""


# def extract_text_from_pdf_content(pdf_content):
#     with fitz.open(stream=pdf_content, filetype="pdf") as doc:
#         text = ""
#         for page in doc:
#             text += page.get_text()
#         return text



def get_file_content(access_token, download_url):
    """Download the file content from the provided URL."""
    headers = {"Authorization": f"Bearer {access_token}"}
    response = requests.get(download_url, headers=headers)
    response.raise_for_status()
    return response.content

def pdf_to_images(pdf_path):
    images = None
    if os.name == "nt":
        # For windows
        print('--- Windows ---')
        images = convert_from_path(pdf_path, poppler_path=r'C:\poppler\poppler-24.08.0\Library\bin')
    else:
        # For Linux
        print('--- Linux ---')
        images = convert_from_path(pdf_path)
    return images

def process_pdf_with_gpt(pdf_path, prompt, client):
    from .common import CommonUtils
    results = []
    images = pdf_to_images(pdf_path)

    for i, image in enumerate(images):
        print(f"Processing page {i + 1}...")
        result = CommonUtils.send_image_to_gpt(client, image, prompt)
        results.append(result)

    return "\n\n".join(results)

def get_pdf_file_content(access_token, download_url, client):
    """Download the file content from the provided URL."""
    PARAMS = {
        "la": "en",
        "hash": "5EA2DA7D1492D105375580EEF2FB088F"
    }

    HEADERS = {
        "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 14_3_1) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.2 Safari/605.1.15",
        "Accept-Language": "en-GB,en;q=0.9,en-US;q=0.8,pt;q=0.7",
        "Authorization": f"Bearer {access_token}"
    }

    CHUNK = 32 * 1024

    with requests.get(download_url, headers=HEADERS, params=PARAMS, stream=True) as response:
        response.raise_for_status()
        with open("temp.pdf", "wb") as output:
            for data in response.iter_content(CHUNK):
                output.write(data)

    prompt = """
    Please analyze the scanned questionnaire page and extract every question along with its corresponding answer(s). For multiple-choice questions, note that there are two types:
  • Radio button questions: Only one option is selected (indicated by a filled radio button).
  • Tick box questions: One or more options may be selected (indicated by one or more ticked boxes).

    Additionally, extract any free-text answers provided on the page.
    
    Format your output exactly as follows:
    Q1: [Text of the question]
    A1: [For radio buttons: the selected option; for tick boxes: a list of all selected options; include any associated free-text answer if present]
    Q2: [Text of the question]
    A2: [Answer(s)]
    ... and so on.
    
    Ensure that your transcription accurately reflects all text from the page.
    """

    all_qna = process_pdf_with_gpt("temp.pdf", prompt, client)
    os.remove("temp.pdf")

    return all_qna

def parse_pdf_content(file_content):
    """Parse the content of a PDF file."""
    with open("temp.pdf", "wb") as f:
        f.write(file_content)
    reader = PdfReader("temp.pdf")
    content = "\n".join([page.extract_text() for page in reader.pages])
    os.remove("temp.pdf")
    return content


def send_to_gpt(client, parsed_content):
    """Send parsed content to GPT for a response."""
    deployment_name_model = config("DEPLOYMENT_NAME")
    prompt = (
        f"I want the response in JSON format. Here is the content: \n{parsed_content}\n"
        f"Please structure the JSON with keys named 'solution_plays' and include in the values the technical capabilities along with a description of each."
        f"Make sure to add all Solution Plays from the content into Json keys"
    )
    response = client.chat.completions.create(
        model=config("MODEL_NAME"),
        # max_tokens=2000,
        response_format={"type": "json_object"},
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content

@log_execution_time
def gpt_response_for_sp(client, prompt):
    deployment_name_model = config("DEPLOYMENT_NAME")
    response = client.chat.completions.create(
        model=config("MODEL_NAME"),
        # max_tokens=350,
        response_format={"type": "json_object"},
        messages=[{"role": "user", "content": prompt}],
    )
    result = response.choices[0].message.content
    return result


def save_response_to_json(data, file_path):
    """Save the GPT response to a JSON file."""
    with open(file_path, "w") as f:
        json.dump(data, f, indent=4)


def set_is_parsed_false(access_token, item_id):
    TAXONOMY_DRIVE_ID = config("TAXONOMY_DRIVE_ID")
    url = f"https://graph.microsoft.com/v1.0/drives/{TAXONOMY_DRIVE_ID}/items/{item_id}/listItem/fields"
    headers = {"Authorization": f"Bearer {access_token}"}
    fields = { "isParsed": "True" }
    response = requests.patch(url,json=fields ,headers=headers)
    response.raise_for_status()
    if response.status_code == "200":
        return True
    return False


def read_json_file(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            data = json.load(file)
        return data
    except FileNotFoundError:
        raise FileNotFoundError(f"The file at {file_path} does not exist.")
    except json.JSONDecodeError as e:
        raise json.JSONDecodeError(f"Error decoding JSON: {e.msg}", e.doc, e.pos)


def taxonomy_processing(client, access_token):
    TAXONOMY_DRIVE_ID = config("TAXONOMY_DRIVE_ID")
    drive_url = f"https://graph.microsoft.com/v1.0/drives/{TAXONOMY_DRIVE_ID}/root/children"
    items = get_sharepoint_items(access_token, drive_url)
    if len(items) == 0:
        return "No Taxonomy file found (or) All files have already been processed!", "", False

    item_id, download_url = get_taxonomy_item_id(access_token, items)
    if item_id == -1:
        return "No Taxonomy file found (or) All files have already been processed!", "", False
    file_content = get_file_content(access_token, download_url)
    set_is_parsed_false(access_token, item_id)

    parsed_content = parse_pdf_content(file_content)
    gpt_response = send_to_gpt(client, parsed_content)

    json_file_path = "response.json"
    save_response_to_json(eval(gpt_response), json_file_path)
    return "File processed and response saved", json_file_path, True



def extract_qna_from_docx(binary_content: bytes) -> dict:
    """
    Extracts questions and answers from a .docx file provided as binary content.

    Args:
        binary_content (bytes): The binary content of the .docx file.


    Returns:
        dict: A dictionary where each question (without "Q:") is a key and the answer (without "A:") is the value.
    """
    # Create a temporary .docx file
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
        temp_filename = temp_file.name
        temp_file.write(binary_content)

    qa_dict = {}
    WORD_NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    try:
        # Open the docx file as a zip and extract XML
        with zipfile.ZipFile(temp_filename) as docx_zip:
            xml_content = docx_zip.read('word/document.xml')

        root = ET.fromstring(xml_content)
        paragraphs = root.findall('.//w:p', WORD_NS)

        current_question = None
        answer_lines = []

        for para in paragraphs:
            text = "".join(node.text for node in para.findall('.//w:t', WORD_NS) if node.text)
            text = text.strip()

            if text.startswith("Q:"):
                if current_question is not None:
                    full_answer = "\n".join(answer_lines).strip()
                    qa_dict[current_question] = full_answer

                current_question = text[len("Q:"):].strip()
                answer_lines = []

            elif text.startswith("A:"):
                answer_text = text[len("A:"):].strip()
                if current_question:
                    answer_lines.append(answer_text)

            else:
                if current_question and text:
                    answer_lines.append(text)

        if current_question is not None:
            full_answer = "\n".join(answer_lines).strip()
            qa_dict[current_question] = full_answer

    finally:
        # Always delete the temp file
        os.remove(temp_filename)

    return qa_dict

def process_docx_content(binary_content: bytes) -> str:
    # Create a temporary .docx file
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
        temp_filename = temp_file.name
        # Write binary content to the temporary .docx file
        temp_file.write(binary_content)

    # Read the content back from the .docx file
    document_content = ''
    try:
        doc = Document(temp_filename)
        for para in doc.paragraphs:
            document_content += para.text + '\n'
    finally:
        # Ensure the temporary file is deleted even if an error occurs
        os.remove(temp_filename)

    # Return the content read from the .docx file
    return document_content.strip()


def get_file_down_url(access_token, items, project_id, delimiter):
    headers = {"Authorization": f"Bearer {access_token}",}
    item_values = items["value"]
    target_ind = None
    time.sleep(5)
    for ind, elem in enumerate(item_values):
        split_name_list =  elem["name"].split(delimiter)
        if len(split_name_list) > 1:
            sp_proj_id = int(re.findall(r"\d+", split_name_list[1])[0]) # ['filename', '70.docx']
            # item_proj_id = int(split_name_list[1])
            if sp_proj_id == int(project_id):
                target_ind = ind
                break

    download_url = item_values[target_ind]["@microsoft.graph.downloadUrl"]
    return download_url

def get_initial_form_by_search(access_token, item_id, client):
    is_pdf = False
    is_docx = False

    init_form_drive=config("INITIAL_FORM_DRIVE")
    url = f"https://graph.microsoft.com/v1.0/drives/{init_form_drive}/items/{item_id}"
    headers = {"Authorization": f"Bearer {access_token}"}
    response = requests.get(url, headers=headers)
    response = response.json()

    download_url = response.get("@microsoft.graph.downloadUrl", None)
    if not download_url:
        raise "There was an issue while getting the Download URL from Sharepoint"

    item_name = response.get("name", None)

    if ".pdf" in item_name:
        is_pdf = True
    elif ".docx" in item_name:
        is_docx = True

    file_content = ""

    if is_pdf:
        print("the Initial form is in PDF Format")
        # Retry mechanism
        max_attempts = 2
        for attempt in range(1, max_attempts + 1):
            try:
                file_content = get_pdf_file_content(access_token, download_url, client)
                break  # Success, exit the loop
            except Exception as e:
                print(f"Attempt {attempt} failed: {e}")
                if attempt == max_attempts:
                    raise  # Re-raise the exception if it's the last attempt
                time.sleep(2 ** attempt)  # Exponential backoff: 2s, 4s, 8s...

        # For testing (check the initial form content by saving in a file locally)
        with open('initial form text.txt', "w") as f:
            f.write(file_content)
            f.close()

    elif is_docx:
        print("the Initial form is is DOCX Format")
        binary_content = get_file_content(access_token, download_url)
        # file_content = process_docx_content(binary_content)
        file_content = str(extract_qna_from_docx(binary_content))

    return file_content

def get_initial_form_content(access_token, project_id):
    time.sleep(5)
    initial_form_drive = config("INITIAL_FORM_DRIVE")
    drive_url = f"https://graph.microsoft.com/v1.0/drives/{initial_form_drive}/root/children"
    #TODO: filter by project id and its download url
    items = get_sharepoint_items(access_token, drive_url)
    if len(items) == 0:
        return "No files found!", False

    download_url = get_file_down_url(access_token, items, project_id, delimiter="_")
    file_content = get_file_content(access_token, download_url)
    return file_content, True

@log_execution_time
def get_discovery_questionnaire(access_token, project_id):
    DISCOVERY_DRIVE = config("DISCOVERY_DRIVE")
    drive_url = f"https://graph.microsoft.com/v1.0/drives/{DISCOVERY_DRIVE}/root/children"
    items = get_sharepoint_items(access_token, drive_url)

    if len(items) == 0:
        return "No files found!", False

    download_url = get_file_down_url(access_token, items, project_id, delimiter="-")
    binary_content = get_file_content(access_token, download_url)
    file_content = process_docx_content(binary_content)
    return file_content, True


def get_discovery_content(access_token, item_id):
    url = f"https://graph.microsoft.com/v1.0/drives/b!g1RPFkGuNkGOxozZZFyUfcWTvdgFKoJFkMbW7oxfQJ7BI2nybhy9Qp-2Uu0XUmby/items/{item_id}"
    headers = {"Authorization": f"Bearer {access_token}"}
    response = requests.get(url, headers=headers)
    response = response.json()
    download_url = response.get("@microsoft.graph.downloadUrl", None)
    if not download_url:
        raise "There was an issue while getting the Download URL from Sharepoint"
    file_content_binary = get_file_content(access_token, download_url)
    questionnaire_content = process_docx_content(binary_content=file_content_binary)

    return questionnaire_content


def get_project_name(access_token, project_id):
    try:
        access_token = get_access_token()
        headers = {
            "Authorization": f"Bearer {access_token}",
            "Content-Type": "application/json"
        }

        site_id = config("SITE_ID")
        project_list_id = config("PROJECT_LIST")
        url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/lists/{project_list_id}/items/{project_id}/fields"

        # PATCH request to update CurrentStep
        response = requests.get(url, headers=headers)
        response_json = response.json()
        project_name = response_json.get("Title")

        if response.status_code != 200:
            raise Exception(f"Failed to get Project List/{project_id} Info: {response.json()}")

        return project_name

    except Exception as e:
        raise Exception(f"Error getting Info for Project-{project_id}: {str(e)}")


def get_template(access_token, template_type):
    """
    Downloads a template file (DOCX or XLSX) based on template_type ('SOW' or 'WBS').
    Preserves formatting and permissions via download URL.
    """

    templates_drive_id = config("TEMPLATES_DRIVE_ID")
    url = f"https://graph.microsoft.com/v1.0/drives/{templates_drive_id}/root/children"
    headers = {"Authorization": f"Bearer {access_token}"}

    response = requests.get(url, headers=headers)
    if response.status_code != 200:
        raise Exception(f"Failed to fetch files: {response.json()}")

    items = response.json().get("value", [])
    target_item = None

    # Find the item that matches the template_type
    for item in items:
        item_id = item["id"]
        item_url = f"https://graph.microsoft.com/v1.0/drives/{templates_drive_id}/items/{item_id}/listItem"
        item_response = requests.get(item_url, headers=headers)
        if item_response.status_code != 200:
            continue

        item_details = item_response.json()
        if item_details["fields"].get("template_type") == template_type:
            target_item = item
            break

    if not target_item:
        raise Exception(f"No template found for type: {template_type}")

    download_url = target_item["@microsoft.graph.downloadUrl"]
    filename = target_item["name"]
    file_ext = os.path.splitext(filename)[1].lower()  # e.g., '.docx' or '.xlsx'

    output_path =  filename

    file_download_response = requests.get(download_url)
    if file_download_response.status_code != 200:
        raise Exception("Failed to download the file.")

    with open(output_path, "wb") as f:
        f.write(file_download_response.content)

    return output_path
